"""
This script reads the output from the image processing workflow:
1. The 'metadata_report.json' file.
2. The folder of timestamped images and videos.

It then generates a Microsoft Word document (.docx) with one page per media file,
containing the image/video frame and its corresponding metadata.
"""
import os
import json
import cv2  # OpenCV is needed to read video frames
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import IMAGE_FOLDER, METADATA_JSON_FILE, OUTPUT_DOCX_FILE, PEOPLE_TO_ADD

def load_metadata_from_json(json_path):
    """
    Loads the structured metadata from the JSON report file.
    """
    if not os.path.exists(json_path):
        print(f"Error: JSON report file not found at '{json_path}'")
        return None
    
    print(f"Loading metadata from '{json_path}'...")
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    print(f"Successfully loaded metadata for {len(data)} photos.")
    return data

def extract_frame_from_video(video_path, output_image_path):
    """
    Extracts the middle frame from a video and saves it as a temporary image file.
    Returns True on success, False on failure.
    """
    try:
        cap = cv2.VideoCapture(video_path)  # pylint: disable=no-member
        if not cap.isOpened():
            print(f"Error: Could not open video file {video_path}")
            return False
        
        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))  # pylint: disable=no-member
        middle_frame_index = total_frames // 2
        
        cap.set(cv2.CAP_PROP_POS_FRAMES, middle_frame_index)  # pylint: disable=no-member
        ret, frame = cap.read()
        cap.release()
        
        if ret:
            cv2.imwrite(output_image_path, frame)  # pylint: disable=no-member
            return True
        return False
    except Exception as e:
        print(f"An error occurred while extracting frame: {e}")
        return False

def create_word_document(media_folder, metadata_dict, people_in_photo):
    """
    Creates a Word document from images and video frames with your specific formatting.
    """
    if not os.path.isdir(media_folder):
        print(f"Error: Media folder not found at '{media_folder}'")
        return

    print("Creating Word document...")
    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for processed_filename, metadata in metadata_dict.items():
        media_path = os.path.join(media_folder, processed_filename)
        
        if not os.path.exists(media_path):
            print(f"Warning: Media file '{processed_filename}' not in folder. Skipping.")
            continue

        print(f"Adding '{metadata.get('OriginalFileName', '')}' to document...")

        original_filename = metadata.get('OriginalFileName', processed_filename)
        is_video = original_filename.lower().endswith(('.mov', '.mp4'))
        path_for_doc = None
        temp_frame_path = None

        if is_video:
            temp_frame_path = os.path.join(media_folder, f"_temp_frame_{processed_filename}.png")
            if extract_frame_from_video(media_path, temp_frame_path):
                path_for_doc = temp_frame_path
            else:
                print(f"Could not extract a frame from video '{processed_filename}'. Skipping.")
                continue
        else:
            path_for_doc = media_path

        if metadata.get('People') == "NULL" and people_in_photo:
            metadata['People'] = people_in_photo

        heading_text = original_filename
        if is_video:
            heading_text += " (Video Frame)"
        heading = document.add_heading(heading_text, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        with Image.open(path_for_doc) as img:
            width_px, height_px = img.size

        if width_px > height_px:  # Landscape
            document.add_picture(path_for_doc, width=Inches(7.0))
        else:  # Portrait or square
            document.add_picture(path_for_doc, height=Inches(5.7))
        
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()
        
        # ==================== MODIFICATION START ====================
        # This new, simplified loop correctly processes all metadata fields.
        
        # A mapping of technical keys to their user-friendly names.
        pretty_names = {
            'GPS_Location': 'GPS Location',
            'Location_Address': 'Location Address',
            'DeviceMake': 'Device Make',
            'DeviceModel': 'Device Model',
            'FocalLength': 'Focal Length',
            'ShutterSpeed': 'Shutter Speed'
        }

        for key, value in metadata.items():
            # Skip fields that are handled manually or not needed in this list
            if key in ['Comments']:
                continue

            # Use the pretty name if available, otherwise use the key as-is
            field_name = pretty_names.get(key, key)
            
            # Your custom capitalization logic
            field_content = value[0].upper() + value[1:] if isinstance(value, str) and value else value

            # Add the formatted metadata line to the document
            p = document.add_paragraph()
            p.add_run(f'{field_name}: ').bold = True
            p.add_run(str(field_content))
            p.paragraph_format.space_after = Pt(0)
        # ===================== MODIFICATION END =====================

        p = document.add_paragraph()
        p.add_run('Description: ').bold = True
        p.paragraph_format.space_after = Pt(0)
        
        comments = metadata.get('Comments')
        if comments and comments != "NULL" and comments != "N/A":
            p.add_run(comments)
        
        document.add_page_break()

        if temp_frame_path and os.path.exists(temp_frame_path):
            os.remove(temp_frame_path)

    try:
        document.save(OUTPUT_DOCX_FILE)
        print(f"\nSuccessfully created Word document: '{OUTPUT_DOCX_FILE}'")
    except Exception as e:
        print(f"Error saving document: {e}")

if __name__ == "__main__":
    if not PEOPLE_TO_ADD:
        print("Error: Please edit the script and set the 'PEOPLE_TO_ADD' variable before running.")
    else:
        loaded_data = load_metadata_from_json(METADATA_JSON_FILE)
        if loaded_data:
            create_word_document(IMAGE_FOLDER, loaded_data, PEOPLE_TO_ADD)

